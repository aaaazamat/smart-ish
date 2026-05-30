"""Word (.docx) rezyumesidan strukturali ma'lumot ajratish.

Gibrid: avval AI (Gemini JSON-mode) bilan, quota tugasa heuristik (zaxira) parse.
Natija FK mapping bilan rezyume yaratishga tayyor dict qaytaradi.

Asosiy funksiya: parse_resume(uploaded_file) -> dict
"""
import json
import logging
import re
from difflib import SequenceMatcher

from django.conf import settings

from .ai_services import _call_gemini, AIServiceError
from .translation_service import _translate_keys
from .models import (
    Profession, Region, Skill, University, UniversityDirection,
)

logger = logging.getLogger(__name__)


# ─────────────────────────────────────────────
# 1) DOCX → matn
# ─────────────────────────────────────────────

def _para_full_text(para) -> str:
    """Paragraf matni + havola (hyperlink) matnini birlashtiradi.

    python-docx'da para.text giperhavola matnini o'z ichiga olmaydi —
    LinkedIn/email havolalarini yo'qotmaslik uchun alohida qo'shamiz.
    """
    parts = [para.text or ""]
    try:
        for hl in para.hyperlinks:
            if hl.text and hl.text not in parts[0]:
                parts.append(hl.text)
    except Exception:
        pass
    return " ".join(p for p in parts if p).strip()


def extract_text_from_docx(uploaded_file) -> str:
    """Yuklangan .docx fayldan MAKSIMAL matnni oladi.

    Qamrab oladi: body paragraflar (+ havolalar), jadvallar (qatorma-qator),
    sarlavha (header) va quyi yozuv (footer). Ro'yxatlar ham para.text orqali keladi.
    Universal — har qanday docx strukturasi (shablon shart emas).
    """
    import docx  # python-docx

    document = docx.Document(uploaded_file)
    lines = []

    # 1) Header (ba'zi rezyumelarda ism/kontakt header'da bo'ladi)
    seen_headers = set()
    for section in document.sections:
        try:
            for para in section.header.paragraphs:
                t = _para_full_text(para)
                if t and t not in seen_headers:
                    seen_headers.add(t)
                    lines.append(t)
        except Exception:
            pass

    # 2) Body — paragraf va jadvallar HUJJAT TARTIBIDA
    #    (python-docx body elementlarini ketma-ket o'qiymiz)
    from docx.table import Table
    from docx.text.paragraph import Paragraph

    body = document.element.body
    for child in body.iterchildren():
        if child.tag.endswith("}p"):
            para = Paragraph(child, document)
            t = _para_full_text(para)
            if t:
                lines.append(t)
        elif child.tag.endswith("}tbl"):
            table = Table(child, document)
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells if c.text.strip()]
                if cells:
                    lines.append("  ".join(dict.fromkeys(cells)))  # takror celllarni olib tashlash

    # 3) Footer
    seen_footers = set()
    for section in document.sections:
        try:
            for para in section.footer.paragraphs:
                t = _para_full_text(para)
                if t and t not in seen_footers:
                    seen_footers.add(t)
                    lines.append(t)
        except Exception:
            pass

    return "\n".join(lines)


# ─────────────────────────────────────────────
# 2) AI parse (Gemini JSON-mode)
# ─────────────────────────────────────────────

_AI_PROMPT = """Sen tajribali HR va rezyume tahlilchisisan. Quyida HAR QANDAY
formatdagi (har xil shablon, til yoki erkin tuzilgan) rezyume matni berilgan.
Undagi ma'lumotni chuqur tushunib, strukturali ajrat va FAQAT JSON qaytar.

MUHIM:
- Matn tartibsiz yoki boshqa shablonda bo'lsa ham, mazmunidan tushunib ajrat.
- Sarlavhalar boshqacha bo'lsa ham (masalan "Work History"="Ish tajribasi",
  "Education"="Ta'lim", "Skills"="Ko'nikmalar") mazmun bo'yicha joylashtir.
- Namuna/placeholder matn ("ISM FAMILIYA", "Lavozim nomi", "masalan...") bo'lsa,
  uni HAQIQIY ma'lumot deb OLMA — bo'sh qoldir.
- Faqat JSON, boshqa hech narsa yozma.

JSON sxemasi (maydon yo'q bo'lsa null yoki bo'sh massiv):
{
  "first_name": "ism",
  "last_name": "familiya",
  "middle_name": "otasining ismi yoki bo'sh",
  "phone_number": "+998XXXXXXXXX yoki null",
  "email": "email yoki null",
  "profession": "kasb/lavozim nomi",
  "profession_detail": "qisqacha ma'lumot / o'zi haqida (2-4 jumla)",
  "career_level": "beginner|junior|middle|fresh_graduate|experienced",
  "expected_salary": null yoki son (so'mda),
  "region": "viloyat/shahar nomi yoki null",
  "skills": ["ko'nikma1", "ko'nikma2"],
  "work_experiences": [
    {"position":"lavozim","organization_name":"kompaniya","start_year":2021,"start_month":1,"end_year":2023,"end_month":12,"is_current":false,"responsibilities":"vazifa/yutuqlar"}
  ],
  "educations": [
    {"degree_level":"secondary_special|bachelor|master|phd","university":"OTM to'liq nomi","direction":"yo'nalish/mutaxassislik","start_year":2017,"end_year":2021,"is_studying":false}
  ],
  "languages": [
    {"language":"uz|ru|en|tr|ko|zh|de|ja|ar","level":"A1|A2|B1|B2|C1|C2"}
  ],
  "certificates": [
    {"name":"sertifikat/kurs nomi","issued_date":"YYYY-MM-DD yoki null"}
  ]
}

Qoidalar:
- Til darajasi: "ona tili/native/родной"→C2, "erkin/свободно/fluent"→C1,
  "yuqori/upper"→B2, "o'rta/intermediate/средний"→B1, "boshlang'ich"→A1,
  A1-C2 ko'rsatilgan bo'lsa o'sha.
- is_current: "Hozirgacha/hozir/present/current" bo'lsa true.
- career_level: tajriba yiliga qarab (0→beginner, 1-2→junior, 3-5→middle,
  5+→experienced, yangi bitiruvchi→fresh_graduate).
- degree_level: "bakalavr/bachelor"→bachelor, "magistr/master"→master,
  "PhD/doktor"→phd, "o'rta maxsus/kollej"→secondary_special.
- Telefonni +998XXXXXXXXX formatiga keltir.
- responsibilities: vazifa va yutuqlarni qisqa, bitta matnda birlashtir.

REZYUME MATNI:
"""


# Parse uchun modellar: avval eng kuchli (sifat), keyin tejamkor (ko'p quota).
PARSE_MODELS = ["gemini-2.5-flash", "gemini-2.0-flash-lite"]


def parse_with_ai(text: str) -> dict:
    """Gemini JSON-mode bilan rezyume matnini strukturalashtirish.

    Eng kuchli model (gemini-2.5-flash) bilan boshlaydi; quota tugasa
    tejamkor modelга (gemini-2.0-flash-lite) o'tadi. Har model uchun barcha
    mavjud kalitlarni sinaydi (round-robin hovuzi).
    """
    keys = _translate_keys() or [getattr(settings, "GEMINI_API_KEY", "")]
    keys = [k for k in keys if k]
    if not keys:
        raise AIServiceError("Gemini kaliti sozlanmagan")

    prompt = _AI_PROMPT + text[:10000]  # uzun rezyume ham sig'sin
    last_err = None

    for model in PARSE_MODELS:
        for api_key in keys:
            try:
                raw = _call_gemini(
                    prompt=prompt,
                    temperature=0.1,
                    max_tokens=4000,
                    response_mime_type="application/json",
                    api_key=api_key,
                    model=model,
                )
                data = json.loads(raw)
                if isinstance(data, dict) and data:
                    logger.info("Resume AI-parse OK (model=%s)", model)
                    return data
            except (AIServiceError, json.JSONDecodeError, ValueError) as e:
                last_err = e
                continue  # keyingi kalit yoki model

    raise AIServiceError(f"AI parse muvaffaqiyatsiz: {last_err}")


# ─────────────────────────────────────────────
# 3) Heuristik (zaxira) parse
# ─────────────────────────────────────────────

# Til nomi → kod
_LANG_NAME_TO_CODE = {
    "o'zbek": "uz", "ozbek": "uz", "узбек": "uz", "uzbek": "uz",
    "rus": "ru", "русск": "ru", "russian": "ru",
    "ingliz": "en", "англ": "en", "english": "en",
    "turk": "tr", "korey": "ko", "xitoy": "zh", "nemis": "de",
    "qoraqalpoq": "kaa", "qozoq": "kk", "tojik": "tg", "arab": "ar",
}

_LEVEL_WORDS = {
    "ona tili": "C2", "native": "C2", "родной": "C2",
    "erkin": "C1", "свободно": "C1", "fluent": "C1", "mukammal": "C1",
    "yuqori": "B2", "o'rta-yuqori": "B2",
    "o'rta": "B1", "intermediate": "B1", "средн": "B1",
    "boshlang'ich": "A1", "beginner": "A1",
}

_SECTION_HEADERS = {
    "summary": ["qisqacha", "о себе", "summary"],
    "experience": ["ish tajribasi", "tajriba", "опыт", "experience"],
    "education": ["ta'lim", "talim", "образование", "education"],
    "skills": ["ko'nikma", "konikma", "навыки", "skills"],
    "languages": ["tillar", "til ", "языки", "languages"],
}


def _detect_level(s: str) -> str:
    low = s.lower()
    for m in re.findall(r"\b([abc][12])\b", low):
        return m.upper()
    for word, lvl in _LEVEL_WORDS.items():
        if word in low:
            return lvl
    return "B1"


def _detect_lang_code(s: str):
    low = s.lower()
    for name, code in _LANG_NAME_TO_CODE.items():
        if name in low:
            return code
    return None


def parse_heuristic(text: str) -> dict:
    """Shablon strukturasiga asoslangan oddiy parse (AI'siz zaxira)."""
    lines = [l.strip() for l in text.split("\n") if l.strip()]
    data = {
        "first_name": "", "last_name": "", "middle_name": "",
        "phone_number": None, "email": None,
        "profession": "", "profession_detail": "",
        "career_level": "junior", "expected_salary": None, "region": None,
        "skills": [], "work_experiences": [], "educations": [], "languages": [],
    }
    if not lines:
        return data

    # 1-qator: ism familiya
    name_parts = lines[0].split()
    if len(name_parts) >= 2:
        data["first_name"] = name_parts[0].title()
        data["last_name"] = name_parts[1].title()
    elif name_parts:
        data["first_name"] = name_parts[0].title()

    # 2-qator: kasb/lavozim (sarlavha bo'lmasa)
    if len(lines) > 1 and not any(h in lines[1].lower() for hs in _SECTION_HEADERS.values() for h in hs):
        data["profession"] = re.sub(r"\(.*?\)", "", lines[1]).strip()

    full = "\n".join(lines)
    # Telefon
    phone = re.search(r"\+?998[\s\-]?\d{2}[\s\-]?\d{3}[\s\-]?\d{2}[\s\-]?\d{2}", full)
    if phone:
        data["phone_number"] = "+" + re.sub(r"[^\d]", "", phone.group())[-12:]
    # Email
    email = re.search(r"[\w.\-]+@[\w\-]+\.\w+", full)
    if email:
        data["email"] = email.group()

    # Bo'limlarga ajratish
    sections = {}
    current = None
    for line in lines:
        low = line.lower()
        matched = None
        for key, headers in _SECTION_HEADERS.items():
            if any(low.startswith(h) or low == h.strip() for h in headers):
                matched = key
                break
        if matched:
            current = matched
            sections[current] = []
        elif current:
            sections.setdefault(current, []).append(line)

    # Summary
    if sections.get("summary"):
        # birinchi izoh qatorini (Masalan:...) o'tkazib yuboramiz
        body = [l for l in sections["summary"] if not l.lower().startswith(("masalan", "2-3 jumla"))]
        data["profession_detail"] = " ".join(body[:3])[:500]

    # Tillar
    for line in sections.get("languages", []):
        code = _detect_lang_code(line)
        if code:
            data["languages"].append({"language": code, "level": _detect_level(line)})

    # Ko'nikmalar
    for line in sections.get("skills", []):
        if line.lower().startswith(("texnik", "kasbiy", "dasturlar", "vositalar")):
            after = line.split(":", 1)[-1] if ":" in line else ""
            for s in re.split(r"[,;]", after):
                s = s.strip()
                if s and len(s) < 50:
                    data["skills"].append(s)
        elif "," in line and len(line) < 200:
            for s in re.split(r"[,;]", line):
                s = s.strip()
                if s and len(s) < 50:
                    data["skills"].append(s)

    return data


# ─────────────────────────────────────────────
# 4) Gibrid parse
# ─────────────────────────────────────────────

def parse_resume(uploaded_file) -> dict:
    """Gibrid: avval AI, muvaffaqiyatsiz bo'lsa heuristik. FK mapping bilan."""
    text = extract_text_from_docx(uploaded_file)
    if not text.strip():
        raise ValueError("Word fayl bo'sh yoki o'qib bo'lmadi")

    parsed = None
    try:
        parsed = parse_with_ai(text)
        logger.info("Resume parsed with AI (%d ta tajriba)", len(parsed.get("work_experiences", [])))
    except (AIServiceError, ValueError, json.JSONDecodeError, KeyError) as e:
        logger.warning("AI parse muvaffaqiyatsiz, heuristikга o'tildi: %s", e)
        parsed = parse_heuristic(text)

    return _apply_fk_mapping(parsed)


# ─────────────────────────────────────────────
# 5) FK mapping
# ─────────────────────────────────────────────

def _best_match(name: str, queryset_values):
    """Eng yaqin nomli yozuv ID sini qaytaradi (oddiy fuzzy)."""
    name_low = (name or "").strip().lower()
    if not name_low:
        return None
    best_id, best_ratio = None, 0.0
    for pk, val in queryset_values:
        ratio = SequenceMatcher(None, name_low, (val or "").lower()).ratio()
        if name_low in (val or "").lower() or (val or "").lower() in name_low:
            ratio = max(ratio, 0.85)
        if ratio > best_ratio:
            best_id, best_ratio = pk, ratio
    return best_id if best_ratio >= 0.6 else None


def match_profession(name):
    return _best_match(name, Profession.objects.values_list("id", "name"))


def match_region(name):
    return _best_match(name, Region.objects.values_list("id", "name"))


def match_skills(names):
    """Ko'nikma nomlarini ID'larga moslaydi; topilmaganini yaratadi."""
    ids = []
    for raw in (names or []):
        nm = (raw or "").strip()
        if not nm or len(nm) > 100:
            continue
        existing = Skill.objects.filter(name__iexact=nm).first()
        if existing:
            ids.append(existing.id)
        else:
            obj, _ = Skill.objects.get_or_create(name=nm, defaults={"name_uz": nm})
            ids.append(obj.id)
    return ids


def match_university(name):
    """OTM nomini ID'ga moslaydi; topilmaganini yaratadi (None agar bo'sh)."""
    nm = (name or "").strip()
    if not nm or len(nm) > 255:
        return None
    pk = _best_match(nm, University.objects.values_list("id", "name"))
    if pk:
        return pk
    obj, _ = University.objects.get_or_create(name=nm, defaults={"name_uz": nm})
    return obj.id


def match_direction(name, university_id):
    """Yo'nalish nomini OTM ichida moslaydi; topilmaganini yaratadi."""
    nm = (name or "").strip()
    if not nm or len(nm) > 255 or not university_id:
        return None
    qs = UniversityDirection.objects.filter(university_id=university_id)
    pk = _best_match(nm, qs.values_list("id", "name"))
    if pk:
        return pk
    obj, _ = UniversityDirection.objects.get_or_create(
        university_id=university_id, name=nm, defaults={"name_uz": nm},
    )
    return obj.id


_VALID_CAREER = {"beginner", "junior", "middle", "fresh_graduate", "experienced"}
_VALID_DEGREE = {"secondary_special", "bachelor", "master", "phd"}
_VALID_LANG = {"uz", "ru", "en", "tr", "ko", "zh", "de", "ja", "hi", "es",
               "fr", "pt", "ur", "id", "kaa", "tg", "kk", "ky", "ar"}
_VALID_LEVEL = {"A1", "A2", "B1", "B2", "C1", "C2"}


def _apply_fk_mapping(parsed: dict) -> dict:
    """Matn nomlarini FK ID / yaroqli choice qiymatlariga aylantiradi."""
    out = dict(parsed)

    # career_level tekshirish
    if out.get("career_level") not in _VALID_CAREER:
        out["career_level"] = "junior"

    # Profession / Region → ID
    out["profession_id"] = match_profession(parsed.get("profession"))
    out["region_id"] = match_region(parsed.get("region"))

    # Skills → ID ro'yxati
    out["skill_ids"] = match_skills(parsed.get("skills"))

    # Languages — yaroqli kod va daraja
    langs = []
    for l in (parsed.get("languages") or []):
        code = l.get("language")
        lvl = (l.get("level") or "").upper()
        if code in _VALID_LANG and lvl in _VALID_LEVEL:
            langs.append({"language": code, "level": lvl})
    out["languages"] = langs

    # Educations — degree_level + university/direction FK mapping
    edus = []
    for e in (parsed.get("educations") or []):
        dl = e.get("degree_level")
        if dl not in _VALID_DEGREE:
            dl = "bachelor"
        uni_id = match_university(e.get("university"))
        dir_id = match_direction(e.get("direction"), uni_id)
        edus.append({
            "degree_level": dl,
            "university_id": uni_id,
            "direction_id": dir_id,
            "start_year": e.get("start_year"),
            "end_year": e.get("end_year"),
            "is_studying": bool(e.get("is_studying")),
        })
    out["educations"] = edus

    # Certificates
    certs = []
    for c in (parsed.get("certificates") or []):
        nm = (c.get("name") or "").strip()
        if nm and len(nm) <= 255:
            certs.append({"name": nm, "issued_date": _normalize_date(c.get("issued_date"))})
    out["certificates"] = certs

    return out


def _normalize_date(raw):
    """'YYYY-MM-DD' yoki 'YYYY' kabi matnni sanaga keltiradi (None agar yaroqsiz)."""
    if not raw:
        return None
    s = str(raw).strip()
    m = re.match(r"(\d{4})-(\d{2})-(\d{2})", s)
    if m:
        return s[:10]
    m = re.match(r"(\d{4})", s)
    if m:
        return f"{m.group(1)}-01-01"  # faqat yil bo'lsa
    return None
